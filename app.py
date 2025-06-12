# Full Streamlit Cover Letter Generator (with fallback input for missing contact info)

import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
import docx
import os
from dotenv import load_dotenv
from datetime import datetime
import re
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
import io

# Load environment variables
load_dotenv()

# Configure page
st.set_page_config(page_title="Cover Letter Generator", page_icon="📝", layout="wide")

# Gemini API setup
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    st.error("❌ GEMINI_API_KEY not found in environment variables!")
    st.stop()

try:
    genai.configure(api_key=api_key)
except Exception as e:
    st.error(f"❌ Error configuring Gemini API: {str(e)}")
    st.stop()

st.title("📝 Cover Letter Generator")
st.markdown("Generate professional cover letters using **Gemini 2.0 Flash**")

# File upload and inputs
cv_file = st.file_uploader("📎 Upload your CV (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

with st.form("form"):
    job_title = st.text_input("Job Title")
    company = st.text_input("Company Name")
    job_desc = st.text_area("Job Description")
    job_reqs = st.text_area("Job Requirements")
    word_len = st.slider("Word Count Target", 40, 800, 100)
    hr_name = st.text_input("HR Name (Optional)")
    hr_role = st.text_input("HR Role (Optional)")
    language = st.radio("Language", ["English", "Bahasa Indonesia"])
    submitted = st.form_submit_button("Generate Cover Letter")

# Helper functions
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_text(file):
    if file.type == "application/pdf":
        return extract_text_from_pdf(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file)
    elif file.type == "text/plain":
        return str(file.read(), "utf-8")
    return ""

def extract_contact_info(text):
    email = re.search(r"[\w\.-]+@[\w\.-]+", text)
    phone = re.search(r"(?<!\d)(\+62|08|62)[\d\s\-]{8,}(?!\d)", text)
    name = next((line.strip() for line in text.splitlines()[:10] if line and line[0].isupper()), None)
    return name, email.group() if email else "", phone.group() if phone else ""

def generate_prompt(cv_text, name, email, phone):
    today = datetime.now().strftime("%d %B %Y")
    hr_line = f"to {hr_name}, {hr_role}" if hr_name and hr_role else hr_name or "the Hiring Manager"
    lang = "Indonesian (Bahasa Indonesia)" if language == "Bahasa Indonesia" else "English"

    return f"""
You are a professional cover letter writer. Your task is to create an engaging, professional, and customized cover letter using the following:

📅 **Date:** {today}

👤 **Applicant Info:**
- Name: {name}
- Email: {email}
- Phone: {phone}

📄 **Resume (analyze for achievements, skills, and experiences):**
{cv_text}

💼 **Job Info:**
- Position: {job_title}
- Company: {company}
- Description: {job_desc}
- Requirements: {job_reqs}

🎯 **Instructions:**
- Language: Write the letter in **{lang}**.
- Length: Target approx. **{word_len} words** (+/- 15%).
- Address to: **{hr_line}**

📝 **Structure & Tone:**
1. **Header:** Applicant's contact, date, and recipient/company details.
2. **Salutation:** Use specific name if given (e.g., "Dear Mr./Ms. X"), or "Dear Hiring Manager".
3. **Intro:** Show enthusiasm and suitability for the role.
4. **Body:**
    - Match top 2–3 job requirements with real achievements/skills from CV.
    - Use real examples and quantify (e.g., "increased efficiency by 20%").
    - Highlight what value you bring to {company}.
5. **Motivation:** Optional — why you want to work at {company}.
6. **Closing:** Reaffirm interest and politely invite follow-up.
7. **Signature:** Full name

    CRITICAL INSTRUCTIONS:
    1. Do not include any placeholder text in square brackets like [Your Name], [Date], [Company Name], [Your Email], [Your Phone], etc. 
    2. Use the actual provided information: {name}, {email}, {phone}, {today_date}, {company}, etc.
    3. Do not include any metadata, instructions, or notes in square brackets in the final output.
    4. The output should be a clean, professional cover letter ready for immediate use.
    5. Remove any text that appears in square brackets [ ] completely from the final output.
    6. Always structure the paragrapgh and text allignment like professional cover letter
    7. Do not include any address and instruction to fill the address

📌 Avoid copying the CV. Instead, synthesize and write a flowing, impactful letter.
"""

def export_pdf(letter_text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    style = ParagraphStyle(name='Justify', parent=styles['Normal'], alignment=TA_JUSTIFY, fontSize=11)
    elements = [Paragraph(p.strip(), style) for p in letter_text.split('\n') if p.strip()]
    doc.build(elements)
    buffer.seek(0)
    return buffer

# Main logic
if submitted and cv_file:
    with st.spinner("Extracting and analyzing CV..."):
        cv_text = extract_text(cv_file)
        name, email, phone = extract_contact_info(cv_text)

        # Manual fallback inputs if missing
        if not name:
            name = st.text_input("⚠️ Name not detected, please input manually:")
        if not email:
            email = st.text_input("⚠️ Email not detected, please input manually:")
        if not phone:
            phone = st.text_input("⚠️ Phone number not detected, please input manually:")

    if name and email and phone:
        with st.spinner("Generating cover letter using Gemini..."):
            prompt = generate_prompt(cv_text, name, email, phone)
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(prompt)
            letter = response.text.strip()

            st.subheader("📄 Generated Cover Letter")
            st.text_area("Preview", letter, height=400)

            pdf = export_pdf(letter)
            st.download_button("📥 Download as PDF", data=pdf, file_name="Cover_Letter.pdf", mime="application/pdf")
    else:
        st.warning("Please complete the missing contact information.")
else:
    st.info("👆 Please fill out the form and upload your CV to generate a cover letter.")






# # Complete Streamlit Cover Letter Generator
# import streamlit as st
# import google.generativeai as genai
# from PyPDF2 import PdfReader
# import docx
# import os
# from dotenv import load_dotenv
# from datetime import datetime
# import re
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.pagesizes import A4
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
# import io

# # Load environment variables
# load_dotenv()

# # Configure page
# st.set_page_config(page_title="Cover Letter Generator", page_icon="📝", layout="wide")

# # Gemini API setup
# api_key = os.getenv("GEMINI_API_KEY")
# if not api_key:
#     st.error("❌ GEMINI_API_KEY not found in environment variables!")
#     st.markdown("""
#     ### 🔧 Setup Required:
    
#     **1. Create a `.env` file in your project root:**
#     ```
#     GEMINI_API_KEY=your_api_key_here
#     ```
    
#     **2. Get your Gemini API Key:**
#     - Visit [Google AI Studio](https://aistudio.google.com/app/apikey)
#     - Create a new API key (free)
#     - Copy and paste it to your `.env` file
    
#     **3. Restart the application**
#     """)
#     st.stop()

# try:
#     genai.configure(api_key=api_key)
# except Exception as e:
#     st.error(f"❌ Error configuring Gemini API: {str(e)}")
#     st.stop()

# # Title and description
# st.title("📝 Cover Letter Generator")
# st.markdown("Generate professional cover letters using **Gemini 2.0 Flash**")

# def extract_text_from_pdf(pdf_file):
#     """Extract text from PDF file"""
#     try:
#         reader = PdfReader(pdf_file)
#         return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
#     except Exception as e:
#         st.error(f"Error reading PDF: {e}")
#         return ""

# def extract_text_from_docx(docx_file):
#     """Extract text from DOCX file"""
#     try:
#         doc = docx.Document(docx_file)
#         return "\n".join(p.text for p in doc.paragraphs)
#     except Exception as e:
#         st.error(f"Error reading DOCX: {e}")
#         return ""

# def extract_text_from_file(uploaded_file):
#     """Extract text from uploaded file based on type"""
#     if uploaded_file.type == "application/pdf":
#         return extract_text_from_pdf(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         return extract_text_from_docx(uploaded_file)
#     elif uploaded_file.type == "text/plain":
#         return str(uploaded_file.read(), "utf-8")
#     else:
#         st.error("Unsupported file type")
#         return ""

# def extract_contact_info(text):
#     """Extract contact information from CV text"""
#     email = re.search(r"[\w\.-]+@[\w\.-]+", text)
#     phone = re.search(r"(?<!\d)(\+62|08|62)[\d\s\-]{8,}(?!\d)", text)
#     name_lines = text.strip().splitlines()[:10]  # Assume name is at the top
#     name = next((line.strip() for line in name_lines if len(line.split()) >= 2 and line[0].isupper()), None)
#     return name, email.group() if email else "", phone.group() if phone else ""

# def clean_metadata(text):
#     """Remove all metadata in square brackets from text"""
#     if not text:
#         return ""
#     # Remove anything between square brackets including the brackets
#     cleaned_text = re.sub(r'\[.*?\]', '', text)
#     # Remove extra whitespace and newlines that might be left
#     cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
#     cleaned_text = re.sub(r'^\s+|\s+$', '', cleaned_text, flags=re.MULTILINE)
#     return cleaned_text.strip()

# def generate_cover_letter(cv_text, job_title, company, job_desc, job_reqs, word_len, name, email, phone, hr_name, hr_role, bahasa):
#     """Generate cover letter using Gemini AI"""
#     today_date = datetime.now().strftime("%d %B %Y")
#     hr_info = f"to {hr_name}, {hr_role}" if hr_name and hr_role else hr_name if hr_name else "the Hiring Manager"
    
#     # Set language instruction
#     language_instruction = "Indonesian (Bahasa Indonesia)" if bahasa == "Bahasa Indonesia" else "English"
    
#     prompt = f"""
#     Anda adalah seorang penulis surat lamaran profesional yang ahli. Tugas Anda adalah membuat surat lamaran yang menarik, profesional, dan sangat disesuaikan berdasarkan informasi yang diberikan.
#      **Gunakan bahasa:** {language_instruction}
#     **Gunakan tanggal hari ini:** {today_date}
    
#     **Informasi Pelamar:**
#     - Nama Lengkap: {name}
#     - Email: {email}
#     - Nomor Telepon: {phone}
    
#     **Konten CV (Resume) Pelamar:**
#     {cv_text}
#     *(Catatan untuk AI: Analisis teks CV ini secara menyeluruh untuk mengidentifikasi pengalaman, keterampilan, pencapaian, dan kualifikasi yang relevan dengan pekerjaan yang dilamar. Jangan hanya menyalin; ekstrak dan sintesis informasi yang paling relevan.)*
    
#     **Informasi Pekerjaan:**
#     - Judul Posisi: {job_title}
#     - Nama Perusahaan: {company}
#     - Deskripsi Pekerjaan: {job_desc}
#     - Persyaratan Pekerjaan: {job_reqs}
    
#     **Informasi Tambahan (Opsional):**
#     - Kepada (Penerima Surat/HR): {hr_info} *(Jika kosong, sapa dengan "Dear Hiring Manager," atau "Yth. Tim Rekrutmen,")*
#     - Perkiraan Panjang Kata: {word_len} *(Targetkan sekitar angka ini, fleksibilitas +/- 15% diperbolehkan)*
    
#     **Panduan Utama Pembuatan Surat Lamaran:**
    
#     1.  **Format Surat Profesional:**
#         *   **Bagian Kepala:**
#             *   Informasi kontak pelamar ({name}, {email}, {phone}) harus jelas di bagian atas.
#             *   Tanggal hari ini ({today_date}).
#             *   Informasi penerima (jika {hr_info} disediakan dan berisi nama spesifik/jabatan, gunakan itu. Jika tidak, cukup nama perusahaan {company} dan alamat jika ada).
#         *   **Salam Pembuka:**
#             *   Sapa penerima secara spesifik jika {hr_info} menyediakan nama (misalnya, "Yth. Bapak [Nama Belakang]," atau "Dear Ms. [Last Name],").
#             *   Jika {hr_info} kosong, umum (misalnya "Yth. Tim Rekrutmen di {company}," atau "Dear Hiring Manager,"). Hindari "To Whom It May Concern" jika memungkinkan.
#         *   **Isi Surat (Konten Inti dan Alur):**
#             *   **Awali surat** dengan menyebutkan posisi ({job_title}) yang dilamar dan di mana Anda melihat lowongan tersebut (jika informasi ini ada atau dapat diasumsikan secara umum, misal "website perusahaan"). Nyatakan antusiasme Anda dan secara singkat mengapa Anda yakin merupakan kandidat yang kuat untuk peran tersebut.
#             *   **Bagian inti surat harus berfokus pada penyesuaian kualifikasi Anda dengan kebutuhan pekerjaan.** Ini adalah bagian krusial.
#                 *   Identifikasi 2-3 persyaratan atau tanggung jawab utama dari {job_reqs} dan {job_desc}.
#                 *   Untuk setiap poin yang diangkat, **tunjukkan, jangan hanya mengatakan.** Jelaskan bagaimana pengalaman, keterampilan, atau pencapaian spesifik dari {cv_text} pelamar secara langsung relevan dan memenuhi persyaratan tersebut. Gunakan contoh konkret dari CV.
#                 *   **Kuantifikasi pencapaian** jika memungkinkan (misalnya, "berhasil meningkatkan efisiensi proses sebesar 15% dalam 6 bulan").
#                 *   Tekankan bagaimana kontribusi pelamar dapat **memberikan nilai tambah** bagi {company}.
#             *   **(Opsional namun sangat dianjurkan jika informasinya ada atau dapat disimpulkan)** Secara singkat, jelaskan **motivasi khusus** pelamar untuk bergabung dengan {company} (misalnya, ketertarikan pada misi perusahaan, nilai-nilai yang dianut, produk inovatif, atau reputasi industri). Anda juga bisa menyinggung bagaimana pelamar melihat dirinya **cocok dengan budaya perusahaan**, jika ada indikasi.
#             *   **Akhiri surat** dengan mengulangi antusiasme Anda untuk posisi tersebut. Sebutkan ketersediaan Anda untuk diskusi lebih lanjut dan sertakan **ajakan bertindak yang sopan dan jelas** (misalnya, "Saya sangat antusias untuk membahas lebih lanjut bagaimana kualifikasi saya dapat mendukung kesuksesan tim Anda di {company}. Terima kasih atas waktu dan pertimbangan Anda.").
#         *   **Salam Penutup:** Gunakan penutup profesional seperti "Hormat saya," atau "Sincerely,".
#         *   **Tanda Tangan:** Nama lengkap pelamar ({name}).
    
#     2.  **Nada dan Gaya:**
#         *   **Profesional dan Antusias:** Nada harus menunjukkan kepercayaan diri, profesionalisme, dan antusiasme yang tulus terhadap peran dan perusahaan.
#         *   **Bahasa yang Jelas dan Ringkas:** Gunakan bahasa yang mudah dipahami, hindari jargon yang tidak perlu kecuali umum dalam industri tersebut. Kalimat harus efektif dan to the point.
#         *   **Proaktif dan Berorientasi pada Solusi:** Bingkai pengalaman sebagai cara Anda memecahkan masalah atau mencapai tujuan.
    
#     3.  **Konten yang Disesuaikan (Sangat Penting!):**
#         *   **Hindari Pernyataan Generik:** Jangan gunakan frasa klise atau pernyataan umum yang bisa berlaku untuk pekerjaan apa pun (misalnya, "Saya pekerja keras," "Saya pembelajar cepat" tanpa bukti pendukung dari CV).
#         *   **Hindari Klaim Berlebihan/Tidak Berdasar:** Semua klaim tentang keterampilan dan pengalaman harus didukung oleh atau dapat disimpulkan dari {cv_text}.
#         *   **Fokus pada Kebutuhan Perusahaan:** Surat lamaran harus menunjukkan pemahaman tentang apa yang dicari {company} (berdasarkan {job_desc} dan {job_reqs}) dan bagaimana pelamar dapat memenuhi kebutuhan tersebut.
    
#     4.  **Detail Teknis:**
#         *   **Gunakan Detail Kontak Nyata:** Pastikan semua detail kontak pelamar ({name}, {email}, {phone}) digunakan secara akurat dan BUKAN placeholder seperti "[Nama Anda]".
#         *   **Panjang Kata:** Usahakan mendekati {word_len} kata. Kualitas dan relevansi lebih penting daripada jumlah kata yang kaku.
#         *   **Tata Bahasa dan Ejaan:** Pastikan bebas dari kesalahan tata bahasa dan ejaan.
    
#     **Output yang Diharapkan:**
#     CRITICAL INSTRUCTIONS:
#     1. Do not include any placeholder text in square brackets like [Your Name], [Date], [Company Name], [Your Email], [Your Phone], etc. 
#     2. Use the actual provided information: {name}, {email}, {phone}, {today_date}, {company}, etc.
#     3. Do not include any metadata, instructions, or notes in square brackets in the final output.
#     4. The output should be a clean, professional cover letter ready for immediate use.
#     5. Remove any text that appears in square brackets [ ] completely from the final output.
#     6. Always structure the paragrapgh and text allignment like professional cover letter
    
#     Output only the complete cover letter text without any additional commentary, metadata, or instructions.
#     """
    
#     model = genai.GenerativeModel("gemini-2.0-flash")
#     response = model.generate_content(prompt)
    
#     # Clean the response from any metadata in square brackets
#     cleaned_response = clean_metadata(response.text)
#     return cleaned_response

# def create_pdf(text):
#     """Create PDF from text with proper formatting"""
#     if not text or not text.strip():
#         st.error("No text provided for PDF creation")
#         return None
        
#     buffer = io.BytesIO()
#     doc = SimpleDocTemplate(buffer, pagesize=A4, 
#                           rightMargin=72, leftMargin=72, 
#                           topMargin=72, bottomMargin=72)
    
#     # Get styles
#     styles = getSampleStyleSheet()
    
#     # Create custom style for body text
#     body_style = ParagraphStyle(
#         'CustomBody',
#         parent=styles['Normal'],
#         fontSize=11,
#         leading=14,
#         alignment=TA_JUSTIFY,
#         spaceAfter=12
#     )
    
#     # Build story
#     story = []
    
#     try:
#         # Clean text from metadata first
#         clean_text = clean_metadata(text)
        
#         # Split into paragraphs - try different approaches
#         paragraphs = []
        
#         # First try splitting by double newlines
#         if '\n\n' in clean_text:
#             paragraphs = [p.strip() for p in clean_text.split('\n\n') if p.strip()]
#         # If no double newlines, split by single newlines but group related lines
#         elif '\n' in clean_text:
#             lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
#             current_para = ""
#             for line in lines:
#                 if line:
#                     if current_para and (line[0].isupper() or len(line) < 50):
#                         # Start new paragraph if line starts with capital or is short
#                         if current_para:
#                             paragraphs.append(current_para.strip())
#                         current_para = line
#                     else:
#                         current_para += " " + line
#             if current_para:
#                 paragraphs.append(current_para.strip())
#         else:
#             # If no newlines, treat as single paragraph
#             paragraphs = [clean_text.strip()]
        
#         # Add paragraphs to story
#         for para in paragraphs:
#             if para and len(para.strip()) > 0:
#                 try:
#                     # Clean paragraph text for reportlab
#                     para_text = para.strip()
#                     # Remove any remaining problematic characters
#                     para_text = para_text.encode('utf-8', 'ignore').decode('utf-8')
                    
#                     # Create paragraph
#                     p = Paragraph(para_text, body_style)
#                     story.append(p)
#                     story.append(Spacer(1, 12))
                    
#                 except Exception as para_error:
#                     st.warning(f"Skipping problematic paragraph: {str(para_error)}")
#                     continue
        
#         # If story is empty, add the raw text
#         if not story:
#             try:
#                 raw_text = clean_text.replace('\n', ' ').strip()
#                 if raw_text:
#                     story.append(Paragraph(raw_text, body_style))
#             except Exception as e:
#                 st.error(f"Failed to add raw text: {str(e)}")
#                 return None
        
#         # Build PDF
#         doc.build(story)
#         buffer.seek(0)
        
#         # Verify buffer has content
#         if buffer.tell() > 0:
#             buffer.seek(0)
#             return buffer
#         else:
#             st.error("Generated PDF is empty")
#             return None
            
#     except Exception as e:
#         st.error(f"Error creating PDF: {str(e)}")
#         return None

# # Sidebar with information
# with st.sidebar:
#     if api_key:
#         st.success("🔑 API Key loaded from .env")
    
#     st.markdown("---")
#     st.header("ℹ️ About")
#     st.markdown("**Powered by Gemini 2.0 Flash**")
#     st.markdown("This tool generates personalized cover letters by analyzing your CV and the job requirements.")
    
#     st.markdown("---")
#     st.header("📋 How to Use")
#     st.markdown("""
#     1. **Upload your CV/Resume**
#     2. **Fill in contact information**
#     3. **Enter job details**
#     4. **Set preferences**
#     5. **Generate cover letter**
#     6. **Download as TXT or PDF**
#     """)
    
#     st.markdown("---")
#     st.header("💡 Tips")
#     st.markdown("""
#     • Use a detailed CV with achievements
#     • Provide complete job description
#     • Include specific requirements
#     • Review and customize output
#     • Check for any placeholder text
#     """)

# # --- MAIN UI STARTS HERE ---
# cv_file = st.file_uploader("📄 Upload your CV (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])

# if cv_file:
#     # Extract text and contact info
#     cv_text = extract_text_from_file(cv_file)
#     name, email, phone = extract_contact_info(cv_text)
    
#     if cv_text:
#         st.success("✅ CV uploaded successfully!")
        
#         # Show extracted text preview
#         with st.expander("📖 Preview CV Content"):
#             st.text_area("Extracted Text", cv_text[:1000] + "..." if len(cv_text) > 1000 else cv_text, height=200, disabled=True)

#     # Contact Information Section
#     st.subheader("👤 Contact Information")
#     col1, col2, col3 = st.columns(3)
    
#     with col1:
#         name = st.text_input("Your Name", value=name if name else "", key="name_input", help="Full name as it appears on your CV")
#     with col2:
#         email = st.text_input("Your Email", value=email if email else "", key="email_input", help="Professional email address")
#     with col3:
#         phone = st.text_input("Your Phone Number", value=phone if phone else "", key="phone_input", help="Contact phone number")

#     # Job Information Section
#     st.subheader("💼 Job Information")
#     job_title = st.text_input("Job Title", key="job_title_input", help="Exact job title from the posting")
#     company = st.text_input("Company Name", key="company_input", help="Full company name")
#     job_desc = st.text_area("Job Description", key="job_desc_input", height=150, help="Complete job description including responsibilities")
#     job_reqs = st.text_area("Job Requirements", key="job_reqs_input", height=150, help="Specific requirements, skills, and qualifications needed")
    
#     # Additional Settings
#     st.subheader("⚙️ Settings")
#     col1, col2 = st.columns(2)
    
#     with col1:
#         word_len = st.slider("Target Word Count", 50, 500, 250, 10, key="word_len_input", help="Approximate length of cover letter")
#         bahasa = st.selectbox("Cover Letter Language", ["English", "Bahasa Indonesia"], key="bahasa_input")
    
#     with col2:
#         hr_name = st.text_input("HR Name (optional)", key="hr_name_input", help="Name of hiring manager or recruiter")
#         hr_role = st.text_input("HR Role (optional)", key="hr_role_input", help="Title of hiring manager or recruiter")

#     # Generate button
#     if st.button("✨ Generate Cover Letter", type="primary", use_container_width=True):
#         # Validate required fields
#         missing_fields = []
#         if not name: missing_fields.append("Name")
#         if not email: missing_fields.append("Email")
#         if not phone: missing_fields.append("Phone")
#         if not job_title: missing_fields.append("Job Title")
#         if not company: missing_fields.append("Company Name")
#         if not job_desc: missing_fields.append("Job Description")
#         if not job_reqs: missing_fields.append("Job Requirements")
        
#         if missing_fields:
#             st.error(f"❗ Please complete the following required fields: {', '.join(missing_fields)}")
#         else:
#             with st.spinner("🤖 Generating your personalized cover letter..."):
#                 try:
#                     result = generate_cover_letter(cv_text, job_title, company, job_desc, job_reqs, word_len, name, email, phone, hr_name, hr_role, bahasa)
                    
#                     # Store in session state
#                     st.session_state.cover_letter = result
#                     st.session_state.company_name = company
#                     st.session_state.applicant_name = name
#                     st.session_state.job_title = job_title
                    
#                     st.success("✅ Cover letter generated successfully!")
                    
#                 except Exception as e:
#                     st.error(f"❌ Error generating cover letter: {str(e)}")

# # Display generated cover letter
# if hasattr(st.session_state, 'cover_letter') and st.session_state.cover_letter:
#     st.markdown("---")
#     st.subheader("📄 Your Cover Letter")
    
#     # Display the cover letter
#     st.text_area("Preview", st.session_state.cover_letter, height=400, key="cover_letter_preview")
    
#     # Statistics
#     words = len(st.session_state.cover_letter.split())
#     chars = len(st.session_state.cover_letter)
#     paragraphs = st.session_state.cover_letter.count('\n\n') + 1
    
#     col_stat1, col_stat2, col_stat3 = st.columns(3)
#     with col_stat1:
#         st.metric("Words", words)
#     with col_stat2:
#         st.metric("Characters", chars)
#     with col_stat3:
#         st.metric("Paragraphs", paragraphs)
    
#     # Download buttons
#     st.subheader("📥 Download Options")
#     col1, col2, col3 = st.columns(3)
    
#     with col1:
#         # TXT Download
#         filename_safe = f"Cover_Letter_{st.session_state.applicant_name.replace(' ', '_')}_{st.session_state.company_name.replace(' ', '_')}.txt"
#         st.download_button(
#             "📄 Download TXT",
#             data=st.session_state.cover_letter,
#             file_name=filename_safe,
#             mime="text/plain",
#             use_container_width=True
#         )
    
#     with col2:
#         # PDF Download
#         try:
#             pdf_buffer = create_pdf(st.session_state.cover_letter)
#             if pdf_buffer:
#                 pdf_data = pdf_buffer.getvalue()
#                 if len(pdf_data) > 1000:  # Check if PDF has substantial content
#                     filename_pdf = f"Cover_Letter_{st.session_state.applicant_name.replace(' ', '_')}_{st.session_state.company_name.replace(' ', '_')}.pdf"
#                     st.download_button(
#                         "📑 Download PDF",
#                         data=pdf_data,
#                         file_name=filename_pdf,
#                         mime="application/pdf",
#                         use_container_width=True
#                     )
#                 else:
#                     st.error("Generated PDF appears to be empty")
#             else:
#                 st.error("Failed to create PDF")
#         except Exception as e:
#             st.error(f"PDF creation error: {str(e)}")
    
#     with col3:
#         # Copy to clipboard helper
#         if st.button("📋 Show for Copy", use_container_width=True):
#             st.code(st.session_state.cover_letter, language=None)
#             st.info("💡 Select the text above and copy it (Ctrl+C / Cmd+C)")
            
#     # Debug information (optional)
#     if st.checkbox("🔍 Show Debug Info"):
#         st.subheader("Debug Information")
#         st.write(f"Cover letter length: {len(st.session_state.cover_letter)} characters")
#         st.write(f"Contains brackets: {'[' in st.session_state.cover_letter}")
#         if '[' in st.session_state.cover_letter:
#             brackets = re.findall(r'\[.*?\]', st.session_state.cover_letter)
#             st.write(f"Brackets found: {brackets}")
        
#         # Show cleaned version
#         cleaned_version = clean_metadata(st.session_state.cover_letter)
#         st.text_area("Cleaned Version Preview", cleaned_version[:500] + "..." if len(cleaned_version) > 500 else cleaned_version, height=150)
        
# else:
#     # Default information when no cover letter is generated
#     st.info("👆 Upload your CV and fill in the job information above to generate a personalized cover letter.")
    
#     st.subheader("✨ What You'll Get")
#     st.markdown("""
#     - **Professional Format**: Proper business letter structure with header, salutation, and closing
#     - **Personalized Content**: Tailored to your CV and the specific job requirements
#     - **Skills Matching**: Highlights your relevant experience and achievements
#     - **Company-Specific**: Shows genuine interest and research about the company
#     - **Multiple Languages**: Support for English and Bahasa Indonesia
#     - **Multiple Formats**: Download as TXT or PDF
#     - **Clean Output**: No placeholder text or metadata
#     """)

# # Footer
# st.markdown("---")
# st.markdown("**💡 Pro Tip**: Always review and customize the generated cover letter before sending. Add specific details about why you're interested in the company and role.")



# import streamlit as st
# import google.generativeai as genai
# from PyPDF2 import PdfReader
# import docx
# import os
# from dotenv import load_dotenv
# from datetime import datetime
# from reportlab.lib.pagesizes import letter, A4
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib.units import inch
# from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
# import io

# # Load environment variables
# load_dotenv()

# # Configure page
# st.set_page_config(
#     page_title="Cover Letter Generator",
#     page_icon="📝",
#     layout="wide"
# )

# # Configure Gemini API
# api_key = os.getenv("GEMINI_API_KEY")

# # Check if API key exists, if not show error and stop
# if not api_key:
#     st.error("❌ GEMINI_API_KEY not found in environment variables!")
#     st.markdown("""
#     ### 🔧 Setup Required:
    
#     **1. Create a `.env` file in your project root:**
#     ```
#     GEMINI_API_KEY=your_api_key_here
#     ```
    
#     **2. Get your Gemini API Key:**
#     - Visit [Google AI Studio](https://aistudio.google.com/app/apikey)
#     - Create a new API key (free)
#     - Copy and paste it to your `.env` file
    
#     **3. Restart the application**
#     """)
#     st.stop()

# try:
#     genai.configure(api_key=api_key)
#     api_configured = True
# except Exception as e:
#     st.error(f"❌ Error configuring Gemini API: {str(e)}")
#     st.stop()

# # Title and description
# st.title("📝 Cover Letter Generator")
# st.markdown("Generate professional cover letters using AI powered by **Gemini 2.0 Flash**")

# # Function to extract text from PDF
# def extract_text_from_pdf(pdf_file):
#     try:
#         pdf_reader = PdfReader(pdf_file)
#         text = ""
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#         return text
#     except Exception as e:
#         st.error(f"Error reading PDF: {str(e)}")
#         return None

# # Function to extract text from DOCX
# def extract_text_from_docx(docx_file):
#     try:
#         doc = docx.Document(docx_file)
#         text = ""
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"
#         return text
#     except Exception as e:
#         st.error(f"Error reading DOCX: {str(e)}")
#         return None

# # Function to extract text from uploaded file
# def extract_text_from_file(uploaded_file):
#     if uploaded_file is not None:
#         file_type = uploaded_file.type
        
#         if file_type == "application/pdf":
#             return extract_text_from_pdf(uploaded_file)
#         elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             return extract_text_from_docx(uploaded_file)
#         elif file_type == "text/plain":
#             return str(uploaded_file.read(), "utf-8")
#         else:
#             st.error("Unsupported file type. Please upload PDF, DOCX, or TXT files.")
#             return None
#     return None

# # Function to create PDF from text
# def create_pdf(text, filename):
#     buffer = io.BytesIO()
#     doc = SimpleDocTemplate(buffer, pagesize=A4, 
#                           rightMargin=72, leftMargin=72,
#                           topMargin=72, bottomMargin=18)
    
#     # Get styles
#     styles = getSampleStyleSheet()
    
#     # Create custom styles
#     title_style = ParagraphStyle(
#         'CustomTitle',
#         parent=styles['Heading1'],
#         fontSize=16,
#         textColor='#333333',
#         spaceAfter=30,
#         alignment=TA_LEFT
#     )
    
#     body_style = ParagraphStyle(
#         'CustomBody',
#         parent=styles['Normal'],
#         fontSize=11,
#         leading=14,
#         textColor='#333333',
#         alignment=TA_JUSTIFY,
#         spaceAfter=12
#     )
    
#     # Build story
#     story = []
    
#     # Split text into paragraphs
#     paragraphs = text.split('\n\n')
    
#     for i, para in enumerate(paragraphs):
#         if para.strip():
#             # Clean the paragraph text for reportlab
#             clean_para = para.strip().replace('\n', ' ')
            
#             # First paragraph gets title style if it looks like a header
#             if i == 0 and (clean_para.startswith('Dear') or len(clean_para) < 100):
#                 story.append(Paragraph(clean_para, body_style))
#             else:
#                 story.append(Paragraph(clean_para, body_style))
            
#             story.append(Spacer(1, 12))
    
#     # Build PDF
#     doc.build(story)
#     buffer.seek(0)
#     return buffer

# # Function to generate cover letter using Gemini 2.0 Flash
# def generate_cover_letter(cv_text, job_description, job_requirements, job_title, company_name, word_length, hr_name=None, hr_role=None):
#     try:
#         # Use Gemini 2.0 Flash model
#         model = genai.GenerativeModel('gemini-2.0-flash')
        
#         # Construct the prompt
#         hr_info = ""
#         if hr_name and hr_role:
#             hr_info = f"Address the cover letter to {hr_name}, {hr_role}."
#         elif hr_name:
#             hr_info = f"Address the cover letter to {hr_name}."
        
#         prompt = f"""
#         You are a professional cover letter writer. Based on the following CV/Resume and job information, create a compelling and personalized cover letter:

#         CV/RESUME CONTENT:
#         {cv_text}

#         JOB INFORMATION:
#         - Job Title: {job_title}
#         - Company Name: {company_name}
#         - Job Description: {job_description}
#         - Job Requirements: {job_requirements}

#         ADDITIONAL INSTRUCTIONS:
#         {hr_info}

#         Always detect my Name,my Phone Number,my Email on CV 

#         Please create a cover letter that:
#         1. Uses proper business letter format with appropriate salutation and closing
#         2. Has a strong opening paragraph that captures attention
#         3. Organizes content in well-structured paragraphs that specifically highlight relevant experience and skills from the CV that match the job requirements
#         4. Directly addresses the specific job requirements mentioned and shows how the candidate meets them
#         5. Shows genuine enthusiasm for the role and demonstrates knowledge about the company
#         6. Includes specific examples and achievements where possible
#         7. Has a compelling closing paragraph with a call to action
#         8. Is professional, engaging, and tailored to the specific role
#         9. Is approximately {word_length} words in length
#         10. Uses active voice and confident language
#         11. Avoids generic statements and clichés
#         12. Maps the candidate's qualifications to the specific requirements listed
#         13. Don't use self claim words like (iam proficient in ),etc..focus on impact on my CV to make it stand out
#         14. Always mention something in my cv that allign with job description or minimun requirement, but don't make it up something that don't exist on my cv

#         Format the cover letter with:
#         - Input my Name,my Phone Number,my Email on cover letter
#         - Date right now
#         - Professional salutation
#         - Well-structured paragraphs (you decide the optimal number and structure)
#         - Professional closing and signature line

#         Make it sound authentic and compelling while maintaining professionalism. Focus especially on how the candidate's background aligns with the specific requirements mentioned. Structure the paragraphs in the most effective way to present the candidate's qualifications.
#         """
        
#         response = model.generate_content(prompt)
#         return response.text
    
#     except Exception as e:
#         st.error(f"Error generating cover letter: {str(e)}")
#         return None

# # Sidebar with information
# with st.sidebar:
#     st.success("🔑 API Key loaded from .env")
        
#     st.markdown("---")
#     st.header("ℹ️ About")
#     st.markdown("**Powered by Gemini 2.0 Flash**")
#     st.markdown("This tool generates personalized cover letters by analyzing your CV and the job requirements.")
    
#     st.markdown("---")
#     st.header("📋 How to Use")
#     st.markdown("""
#     1. **Upload your CV/Resume**
#     2. **Enter job title & company**
#     3. **Add job description & requirements**
#     4. **Set word count preference**
#     5. **Add HR info** (optional)
#     6. **Click Generate**
#     7. **Download as TXT or PDF**
#     """)
    
#     st.markdown("---")
#     st.header("💝 Tips")
#     st.markdown("""
#     • Use a detailed CV with specific achievements
#     • Provide complete job description
#     • List specific requirements separately
#     • Adjust word count based on company preference
#     • Include technical skills and qualifications
#     • Add company-specific information
#     • Review and customize the output
#     """)

# # Create two columns
# col1, col2 = st.columns([1, 1])

# with col1:
#     st.header("📋 Input Information")
    
#     # File uploader for CV/Resume
#     uploaded_file = st.file_uploader(
#         "Upload your CV/Resume",
#         type=['pdf', 'docx', 'txt'],
#         help="Supported formats: PDF, DOCX, TXT"
#     )
    
#     # Extract text from uploaded file
#     cv_text = ""
#     if uploaded_file:
#         with st.spinner("Extracting text from file..."):
#             cv_text = extract_text_from_file(uploaded_file)
#             if cv_text:
#                 st.success("✅ CV/Resume uploaded successfully!")
#                 with st.expander("Preview extracted text"):
#                     st.text_area("Extracted Content", cv_text[:500] + "..." if len(cv_text) > 500 else cv_text, height=150, disabled=True)
    
#     # Job information inputs
#     st.subheader("Job Information")
#     job_title = st.text_input(
#         "Job Title/Role *", 
#         placeholder="e.g., Senior Software Engineer",
#         help="Enter the exact job title from the job posting"
#     )
#     company_name = st.text_input(
#         "Company Name *", 
#         placeholder="e.g., Google Indonesia",
#         help="Enter the full company name"
#     )
#     job_description = st.text_area(
#         "Job Description *", 
#         placeholder="Paste the complete job description here including responsibilities and qualifications...",
#         height=150,
#         help="Include the full job posting for best results"
#     )
    
#     job_requirements = st.text_area(
#         "Job Requirements *", 
#         placeholder="List the specific requirements, skills, qualifications, and experience needed for this role...",
#         height=150,
#         help="Include technical skills, years of experience, education requirements, certifications, etc."
#     )
    
#     # Optional HR information
#     st.subheader("HR Information (Optional)")
#     col_hr1, col_hr2 = st.columns(2)
#     with col_hr1:
#         hr_name = st.text_input("HR Name", placeholder="e.g., Sarah Johnson")
#     with col_hr2:
#         hr_role = st.text_input("HR Role", placeholder="e.g., HR Manager")
    
#     # Word length control
#     st.subheader("Cover Letter Settings")
#     word_length = st.slider(
#         "Target Word Count", 
#         min_value=20, 
#         max_value=600, 
#         value=350, 
#         step=10,
#         help="Choose the approximate length of your cover letter"
#     )
    
#     # Generate button
#     generate_btn = st.button("🚀 Generate Cover Letter", type="primary", use_container_width=True)

# with col2:
#     st.header("📄 Generated Cover Letter")
    
#     # Check if all required fields are filled
#     if generate_btn:
#         if not cv_text:
#             st.error("❌ Please upload your CV/Resume first!")
#         elif not job_title:
#             st.error("❌ Please enter the job title!")
#         elif not company_name:
#             st.error("❌ Please enter the company name!")
#         elif not job_description:
#             st.error("❌ Please enter the job description!")
#         elif not job_requirements:
#             st.error("❌ Please enter the job requirements!")
#         else:
#             with st.spinner("🤖 AI is crafting your personalized cover letter..."):
#                 cover_letter = generate_cover_letter(
#                     cv_text, job_description, job_requirements, job_title, company_name, word_length, hr_name, hr_role
#                 )
                
#                 if cover_letter:
#                     st.success("✅ Cover letter generated successfully!")
                    
#                     # Store in session state for persistence
#                     st.session_state.cover_letter = cover_letter
#                     st.session_state.company_name = company_name
#                     st.session_state.job_title = job_title

#     # Display cover letter if it exists in session state
#     if 'cover_letter' in st.session_state:
#         cover_letter = st.session_state.cover_letter
#         company_name = st.session_state.company_name
#         job_title = st.session_state.job_title
        
#         # Display the cover letter
#         st.text_area(
#             "Your Cover Letter:",
#             cover_letter,
#             height=600,
#             key="cover_letter_output"
#         )
        
#         # Action buttons
#         col_btn1, col_btn2, col_btn3 = st.columns(3)
        
#         with col_btn1:
#             # Download TXT button
#             st.download_button(
#                 label="📥 Download TXT",
#                 data=cover_letter,
#                 file_name=f"cover_letter_{company_name.replace(' ', '_')}_{job_title.replace(' ', '_')}.txt",
#                 mime="text/plain",
#                 use_container_width=True
#             )
        
#         with col_btn2:
#             # Download PDF button
#             try:
#                 pdf_buffer = create_pdf(cover_letter, f"cover_letter_{company_name}_{job_title}")
#                 st.download_button(
#                     label="📑 Download PDF",
#                     data=pdf_buffer,
#                     file_name=f"cover_letter_{company_name.replace(' ', '_')}_{job_title.replace(' ', '_')}.pdf",
#                     mime="application/pdf",
#                     use_container_width=True
#                 )
#             except Exception as e:
#                 st.error(f"PDF generation error: {str(e)}")
#                 st.info("💡 Install reportlab: `pip install reportlab`")
        
#         with col_btn3:
#             # Copy to clipboard info
#             if st.button("📋 Show for Copy", use_container_width=True):
#                 st.code(cover_letter, language=None)
#                 st.info("💡 Select the text above and copy it (Ctrl+C / Cmd+C)")
        
#         # Additional options
#         st.markdown("---")
#         st.subheader("📊 Document Statistics")
#         words = len(cover_letter.split())
#         chars = len(cover_letter)
#         col_stat1, col_stat2, col_stat3 = st.columns(3)
#         with col_stat1:
#             st.metric("Words", words)
#         with col_stat2:
#             st.metric("Characters", chars)
#         with col_stat3:
#             st.metric("Paragraphs", cover_letter.count('\n\n') + 1)
#     else:
#         # Default view when no generation is done
#         st.info("👈 Fill in the information on the left and click 'Generate Cover Letter' to create your personalized cover letter.")
        
#         # Show example preview
#         st.subheader("📄 Sample Output Preview")
#         st.markdown("""
#         Your generated cover letter will appear here with:
        
#         ✅ **Professional format** with proper business letter structure  
#         ✅ **Personalized content** based on your CV and job requirements  
#         ✅ **Compelling opening** that captures attention  
#         ✅ **Relevant experience** highlighting your best qualifications  
#         ✅ **Requirements mapping** showing how you meet specific job requirements  
#         ✅ **Company-specific** details showing genuine interest  
#         ✅ **Strong closing** with clear call to action  
#         """)

